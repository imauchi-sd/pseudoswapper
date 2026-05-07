"""Tests for .docx document mode support (Stage 1)."""
from pathlib import Path

import pytest
import docx

from pseudoswapper.modes.document import redact_document
from pseudoswapper.session import clear_session
from tests.conftest import make_config

FIXTURE = Path("tests/fixtures/sample_document.docx")


@pytest.fixture()
def cfg():
    return make_config(
        company_terms=["Acme Corporation", "Acme Corp", "Project Nightingale", "acme.com"],
        employees=[
            {"full_name": "John Doe", "email": "john.doe@acme.com", "username": "jdoe"},
            {"full_name": "Jane Smith", "email": "j.smith@acme.com", "username": "jsmith"},
        ],
    )


@pytest.fixture(autouse=True)
def clean_session(tmp_path):
    yield
    clear_session(tmp_path)


def _all_text(path: Path) -> str:
    """Extract all paragraph text from a docx file."""
    doc = docx.Document(str(path))
    return "\n".join(
        "".join(r.text for r in p.runs)
        for p in doc.paragraphs
    )


# ── Output file characteristics ──────────────────────────────────────────────

def test_output_suffix_is_redacted_docx(cfg, tmp_path):
    import shutil
    src = tmp_path / "report.docx"
    shutil.copy(FIXTURE, src)
    out = redact_document(src, cfg, tmp_path)
    assert out.name == "report.redacted.docx"
    assert out.exists()


def test_output_is_valid_docx(cfg, tmp_path):
    import shutil
    src = tmp_path / "report.docx"
    shutil.copy(FIXTURE, src)
    out = redact_document(src, cfg, tmp_path)
    # Should not raise
    docx.Document(str(out))


def test_output_written_alongside_input(cfg, tmp_path):
    import shutil
    src = tmp_path / "memo.docx"
    shutil.copy(FIXTURE, src)
    out = redact_document(src, cfg, tmp_path)
    assert out.parent == src.parent


# ── PII removal ───────────────────────────────────────────────────────────────

def test_employee_name_replaced(cfg, tmp_path):
    src = tmp_path / "doc.docx"
    doc = docx.Document()
    doc.add_paragraph("The owner is John Doe.")
    doc.save(str(src))
    out = redact_document(src, cfg, tmp_path)
    content = _all_text(out)
    assert "John Doe" not in content
    assert "[PERSON_" in content


def test_company_term_replaced(cfg, tmp_path):
    src = tmp_path / "doc.docx"
    doc = docx.Document()
    doc.add_paragraph("This project is run by Acme Corporation.")
    doc.save(str(src))
    out = redact_document(src, cfg, tmp_path)
    content = _all_text(out)
    assert "Acme Corporation" not in content
    assert "[COMPANY_" in content


def test_email_replaced(cfg, tmp_path):
    src = tmp_path / "doc.docx"
    doc = docx.Document()
    doc.add_paragraph("Contact us at hello@example.com for details.")
    doc.save(str(src))
    out = redact_document(src, cfg, tmp_path)
    content = _all_text(out)
    assert "hello@example.com" not in content
    assert "[EMAIL_" in content


def test_token_consistent_across_paragraphs(cfg, tmp_path):
    src = tmp_path / "doc.docx"
    doc = docx.Document()
    doc.add_paragraph("John Doe called.")
    doc.add_paragraph("John Doe left a voicemail.")
    doc.save(str(src))
    out = redact_document(src, cfg, tmp_path)
    content = _all_text(out)
    assert content.count("[PERSON_1]") == 2


def test_two_employees_get_distinct_tokens(cfg, tmp_path):
    src = tmp_path / "doc.docx"
    doc = docx.Document()
    doc.add_paragraph("John Doe and Jane Smith attended the meeting.")
    doc.save(str(src))
    out = redact_document(src, cfg, tmp_path)
    content = _all_text(out)
    assert "[PERSON_1]" in content
    assert "[PERSON_2]" in content


# ── Run-split handling ────────────────────────────────────────────────────────

def test_run_split_name_detected(cfg, tmp_path):
    """'John' and 'Doe' in separate runs must still be replaced as 'John Doe'."""
    src = tmp_path / "split.docx"
    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run("John ")
    r = p.add_run("Doe")
    r.bold = True
    p.add_run(" sent the report.")
    doc.save(str(src))
    out = redact_document(src, cfg, tmp_path)
    content = _all_text(out)
    assert "John Doe" not in content
    assert "[PERSON_" in content


# ── Passthrough ───────────────────────────────────────────────────────────────

def test_passthrough_ip_preserved(cfg, tmp_path):
    src = tmp_path / "doc.docx"
    doc = docx.Document()
    doc.add_paragraph("Alert from 192.168.1.1 involving John Doe.")
    doc.save(str(src))
    out = redact_document(src, cfg, tmp_path, passthrough_types={"IP"})
    content = _all_text(out)
    assert "192.168.1.1" in content
    assert "John Doe" not in content


# ── Full fixture ──────────────────────────────────────────────────────────────

def test_full_fixture_contains_no_known_pii(cfg, tmp_path):
    import shutil
    src = tmp_path / "sample_document.docx"
    shutil.copy(FIXTURE, src)
    out = redact_document(src, cfg, tmp_path)
    content = _all_text(out)
    assert "John Doe" not in content
    assert "Jane Smith" not in content
    assert "Acme Corporation" not in content
    assert "Project Nightingale" not in content
    assert "john.doe@acme.com" not in content
