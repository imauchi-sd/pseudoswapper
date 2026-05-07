from __future__ import annotations

from pathlib import Path

import docx
from docx.oxml.ns import qn


def _iter_paragraphs(doc: docx.Document):
    """Yield every Paragraph in the document: body, tables, headers, footers."""
    yield from doc.paragraphs

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs

    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                yield from hf.paragraphs


def _para_text(para) -> str:
    return "".join(run.text for run in para.runs)


def _set_para_text(para, text: str) -> None:
    """Replace all runs in *para* with a single run containing *text*.

    Paragraph-level formatting (style, spacing, indentation) is preserved.
    Intra-run formatting (bold, italic on individual words) is intentionally
    discarded — the output is consumed by an AI, not a human reader.
    """
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def extract_text(path: Path) -> str:
    """Return all paragraph text joined by newlines, suitable for PII detection."""
    doc = docx.Document(str(path))
    lines: list[str] = []
    for para in _iter_paragraphs(doc):
        lines.append(_para_text(para))
    return "\n".join(lines)


def apply_token_map(src: Path, token_map: dict[str, str], out: Path) -> None:
    """Write a redacted copy of *src* to *out* with tokens applied paragraph-by-paragraph."""
    if not token_map:
        import shutil
        shutil.copy2(src, out)
        return

    import re
    sorted_keys = sorted(token_map, key=len, reverse=True)
    pattern = re.compile("|".join(re.escape(k) for k in sorted_keys))

    doc = docx.Document(str(src))
    for para in _iter_paragraphs(doc):
        original = _para_text(para)
        replaced = pattern.sub(lambda m: token_map[m.group(0)], original)
        if replaced != original:
            _set_para_text(para, replaced)

    doc.save(str(out))
